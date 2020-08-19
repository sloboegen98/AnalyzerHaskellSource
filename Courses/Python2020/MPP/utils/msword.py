from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENTATION

from urllib.request import urlopen
from io import BytesIO

import requests

# Sorry for this function
def update_link(href):
    xs = href.split('/')
    corr_href = ''
    for i in range(len(xs)):
        if i != 3 and i < 7:
            corr_href += xs[i] + '/'
    
    return corr_href[:-1]


class DocxExecutor:
    
    def __init__(self, doc_name='new.docx'):
        self.document = Document()
        self.name = doc_name

    def create_page(self):
        self.document.add_page_break()

    def add_heading(self, heading):
        self.document.add_heading(heading, level=0)

    def add_subheading(self, subheading):
        self.document.add_heading(subheading, level=1)

    def add_image(self, name):
        if 'http' in name or 'upload' in name: 
            self._add_imageNet(name)
        else:
            self._add_imageLoc(name)
        
        self.save_doc()

    def _add_imageLoc(self, name):
        self.document.add_picture(name)

    def _add_imageNet(self, name):
        if 'http' not in name:
            name = 'https://' + name
        response = requests.get(name)
        binary_img = BytesIO(response.content)
        try:
            self.document.add_picture(binary_img, width=Inches(2))
        except:
            name = name.replace('400px', '800px')
            response = requests.get(name)
            binary_img = BytesIO(response.content)
            try:
                self.document.add_picture(binary_img, width=Inches(2))
            except:
                print('Picture {} load error...'. format(name))

    def add_text(self, text):
        self.document.add_paragraph(text)

    def set_name(self, text):
        self.name = text

    def set_orientation(self, mod):
        if mod == 'L':
            self.document.sections[-1].orientation = WD_ORIENTATION.LANDSCAPE 
        elif mod == 'P':
            self.document.sections[-1].orientation = WD_ORIENTATION.PORTRAIT
            
    def save_doc(self):
        self.document.save(self.name)
